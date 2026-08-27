#!/usr/bin/env python3
"""
BPS Report Generator — Laporan Penilaian Biopsikososial
Jana report .docx dan .pdf terus dari borang web.
Jalan: python app.py, then buka http://localhost:5000
"""

import os
import sys
import json
import secrets
import requests
from datetime import datetime
from packaging import version

def _resource_path(rel=''):
    """Cari fail sumber bila jalan sbg .py biasa ATAU .exe (PyInstaller _MEIPASS).

    Dalam EXE, fail (templates/, static/) dikumpulkan ke folder _MEIPASS
    yang read-only. Fail boleh-tulis (output/, config.txt) pula letak
    kat folder yang sama dengan exe supaya kekal selepas tutup.
    """
    if getattr(sys, '_MEIPASS', ''):
        base = sys._MEIPASS
    else:
        base = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base, rel) if rel else base

from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
# WeasyPrint optional — fallback to print-friendly HTML
try:
    from weasyprint import HTML
    HAVE_WEASYPRINT = True
except (ImportError, OSError):
    HAVE_WEASYPRINT = False

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('BPS_SECRET_KEY', secrets.token_hex(32))

# ── Versi App & Update ──
# Naikkan APP_VERSION bila ada perubahan. Update diagihkan guna manifest.json
# (lihat fungsi /check_update dan /apply_update di bawah).
APP_VERSION = '1.2.5'

# Flag: betul ke app ni jalan sebagai EXE PyInstaller?
# Dalam EXE, auto-update dimatikan (fail sumber read-only dalam _MEIPASS).
IS_EXE = bool(getattr(sys, '_MEIPASS', ''))

# ── Lokasi Output (Dropbox / config / local) ──
APP_DIR = _resource_path('')
# Folder "boleh tulis" = sebelah exe bila dalam EXE, else APP_DIR.
if getattr(sys, '_MEIPASS', ''):
    EXE_DIR = os.path.dirname(os.path.abspath(sys.executable))
    LOCAL_OUTPUT = os.path.join(EXE_DIR, 'output')
else:
    EXE_DIR = APP_DIR
    LOCAL_OUTPUT = os.path.join(APP_DIR, 'output')
os.makedirs(LOCAL_OUTPUT, exist_ok=True)

def _baca_config():
    """Baca OUTPUT_PATH dari config.txt (jika ada)."""
    cfg_path = os.path.join(EXE_DIR, 'config.txt')
    if os.path.exists(cfg_path):
        try:
            with open(cfg_path, encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line.startswith('OUTPUT_PATH'):
                        val = line.split('=', 1)[1].strip().strip('"').strip("'")
                        if val:
                            return val
        except Exception:
            pass
    return ''

def _baca_config_value(key):
    """Baca sebarang KEY = value dari config.txt (contoh: UPDATE_URL).
    Guna LAST match supaya edit/manual override di mana-mana baris berkesan."""
    cfg_path = os.path.join(EXE_DIR, 'config.txt')
    val = ''
    if os.path.exists(cfg_path):
        try:
            with open(cfg_path, encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line.startswith(key + '=') or line.startswith(key + ' ='):
                        v = line.split('=', 1)[1].strip().strip('"').strip("'")
                        if v:
                            val = v
        except Exception:
            pass
    return val

# URL manifest update — letak di config.txt:
#   UPDATE_URL = https://raw.githubusercontent.com/<user>/<repo>/main/update_manifest.json
# Kalau kosong, butang Semak Update akan tunjuk mesej "belum dikonfigurasi".
UPDATE_URL = _baca_config_value('UPDATE_URL')

# Maklumat hospital & pegawai (untuk surat/dokumen) — kawan boleh tukar di config.txt
NAMA_HOSPITAL = _baca_config_value('NAMA_HOSPITAL') or 'Hospital Hulu Terengganu'
NAMA_PEGAWAI = _baca_config_value('NAMA_PEGAWAI') or 'NAMA PEGAWAI'
NAMA_SISTEM  = _baca_config_value('NAMA_SISTEM') or 'Kementerian Kesihatan Malaysia'

def _resolve_output_base():
    """Tentukan folder output asas mengikut keutamaan:
    1) Path dalam config.txt (jika ada & wujud)
    2) Dropbox KKM (jika folder wujud)
    3) Folder output/ tempatan
    """
    # 1) Config
    cfg = _baca_config()
    if cfg:
        cfg = os.path.expanduser(cfg)
        if os.path.isdir(cfg):
            return cfg
        # Config ada tapi folder tak wujud — cuba buat
        try:
            os.makedirs(cfg, exist_ok=True)
            return cfg
        except Exception:
            pass

    # 2) Dropbox
    dropbox = os.path.expanduser('D:/Dropbox/01. Medical Social Work HHT')
    if os.path.isdir(dropbox):
        return dropbox

    # 3) Local
    return LOCAL_OUTPUT

OUTPUT_BASE = _resolve_output_base()

# Folder mengikut jenis
DROPBOX_BPS = os.path.join(OUTPUT_BASE, f'Laporan BPS {datetime.now().year}')
DROPBOX_TDI = os.path.join(OUTPUT_BASE, '5. Tabung Darul Iman (TDI)', f'TDI {datetime.now().year}')
DROPBOX_APP = os.path.join(OUTPUT_BASE, 'BPS App Data')

# Simpan OUTPUT_DIR untuk compat
OUTPUT_DIR = DROPBOX_APP if OUTPUT_BASE == LOCAL_OUTPUT else OUTPUT_BASE

# Pastikan folder wujud
for d in [DROPBOX_BPS, DROPBOX_TDI, DROPBOX_APP]:
    try:
        os.makedirs(d, exist_ok=True)
    except Exception:
        pass

# ── OpenRouter AI Key ──
# Baca dari .env Hermes
HERMES_ENV = os.path.expanduser('~/AppData/Local/hermes/.env')
OPENROUTER_API_KEY = ''
if os.path.exists(HERMES_ENV):
    with open(HERMES_ENV) as f:
        for line in f:
            line = line.strip()
            if line.startswith('OPENROUTER_API_KEY=') and '***' not in line:
                OPENROUTER_API_KEY = line.split('=', 1)[1].strip()
                break
AI_MODEL = 'deepseek/deepseek-v4-flash'

# ─────────────────────────────────────────────────
# Data rujukan (dropdown)
# ─────────────────────────────────────────────────
BANGSA = ['Melayu', 'Cina', 'India', 'Bumiputera Sabah', 'Bumiputera Sarawak', 'Lain-lain']
AGAMA = ['Islam', 'Buddha', 'Hindu', 'Kristian', 'Lain-lain']
JANTINA = ['Lelaki', 'Perempuan']
STATUS_KAHWIN = ['Bujang', 'Kahwin', 'Duda/Janda', 'Berpisah']
STATUS_PENDIDIKAN = ['Tidak Bersekolah', 'Sekolah Rendah', 'Sekolah Menengah', 'STPM/Diploma', 'Ijazah', 'Lain-lain']
KEWARGANEGARAAN = ['Warganegara', 'Bukan Warganegara']
JENIS_KEDIAMAN = ['Rumah Banglo', 'Rumah Teres', 'Kondominium', 'Flat', 'Rumah Kedai',
                  'Rumah Panjang', 'Rumah Kampung', 'Projek Perumahan Rakyat (PPR)',
                  'Rumah Setinggan', 'Lain-lain']
STATUS_KEDIAMAN = ['Sendiri/Milik Keluarga', 'Sewa', 'Sewa Bilik', 'Majikan', 'Menumpang', 'RAKR/Institusi']
UTILITI = ['Air', 'Elektrik', 'Tandas', 'Telefon/Handphone', 'Internet/Wifi', 'TV', 'Peti Ais', 'Dapur (Gas/Elektrik/Kayu)']
AGENSI_BANTUAN = ['JKM', 'MAIDAM', 'Yayasan Kebajikan Negara', 'Tabung Bencana',
                  'Zakat', 'Lembaga Hasil Dalam Negeri', 'Lain-lain']

# ─────────────────────────────────────────────────
# Laman utama — borang
# ─────────────────────────────────────────────────
@app.route('/')
def index():
    return render_template('index.html',
        bangsa=BANGSA, agama=AGAMA, jantina=JANTINA,
        status_kahwin=STATUS_KAHWIN, status_pendidikan=STATUS_PENDIDIKAN,
        kewarganegaraan=KEWARGANEGARAAN, jenis_kediaman=JENIS_KEDIAMAN,
        status_kediaman=STATUS_KEDIAMAN, utiliti=UTILITI,
        agensi_bantuan=AGENSI_BANTUAN,
        today=datetime.now().strftime('%d/%m/%Y'),
        has_ai=bool(OPENROUTER_API_KEY),
        output_path=OUTPUT_BASE,
        app_version=APP_VERSION,
        is_dropbox=('Dropbox' in OUTPUT_BASE),
        nama_sistem=NAMA_SISTEM)

# ─────────────────────────────────────────────────
# Jana laporan
# ─────────────────────────────────────────────────
@app.route('/generate', methods=['POST'])
def generate():
    data = request.form
    fmt = data.get('format', 'docx')  # docx / pdf

    # Validasi IC
    ic_raw = data.get('ic', '').strip()
    valid_ic, ic_error, _ = _validate_ic(ic_raw)
    if ic_raw and not valid_ic:
        return jsonify({'error': ic_error}), 400

    if fmt == 'docx':
        return generate_docx(data)
    else:
        return generate_pdf(data)


def _validate_ic(ic_raw):
    """Validate No. Kad Pengenalan Malaysia.
    Format: 12 digit numeric (YYMMDD-BB-NNNN, dash optional).
    Return: (is_valid: bool, error_message: str, cleaned_12digit: str)
    """
    if not ic_raw:
        return False, 'No. Kad Pengenalan kosong.', ''

    # Strip semua non-digit
    digits = ''.join(ch for ch in ic_raw if ch.isdigit())
    if len(digits) != 12:
        return False, f'No. Kad Pengenalan mesti 12 digit (anda masukkan {len(digits)} digit).', digits

    yy = int(digits[0:2])
    mm = int(digits[2:4])
    dd = int(digits[4:6])
    place = int(digits[6:8])
    serial = digits[8:12]

    # Bulan mesti 01-12
    if mm < 1 or mm > 12:
        return False, f'Bulan dalam IC tidak sah ({mm:02d}).', digits

    # Hari mesti 01-31 (ikut bulan)
    if dd < 1 or dd > 31:
        return False, f'Hari dalam IC tidak sah ({dd:02d}).', digits

    # Pastikan tarikh wujud dalam kalendar
    full_year = 2000 + yy if yy <= 24 else 1900 + yy
    try:
        from datetime import date
        date(full_year, mm, dd)
    except ValueError:
        return False, f'Tarikh lahir dalam IC tidak sah ({dd:02d}/{mm:02d}/{full_year}).', digits

    # Kod negeri BB — 01-59 umum; 00 & 70-99 kecil kemungkinan
    if place < 1 or place > 59:
        return False, f'Kod negeri tempat lahir dalam IC tidak sah ({place:02d}).', digits

    # Serial number 4 digit sudah pasti oleh slicing
    if not serial.isdigit():
        return False, 'Nombor siri akhir IC tidak sah.', digits

    return True, '', digits


def prepare_data(d):
    """Kumpul semua data dari form untuk template."""
    return {
        # 1. Maklumat Pesakit
        'rujukan_fail': d.get('rujukan_fail', ''),
        'tarikh_daftar': d.get('tarikh_daftar', datetime.now().strftime('%d/%m/%Y')),
        'nama': d.get('nama', ''),
        'ic': d.get('ic', ''),
        'alamat': d.get('alamat', ''),
        'telefon': d.get('telefon', ''),
        'tarikh_lahir': d.get('tarikh_lahir', ''),
        'umur': d.get('umur', ''),
        'tempat_lahir': d.get('tempat_lahir', ''),
        'jantina': d.get('jantina', ''),
        'bangsa': d.get('bangsa', ''),
        'agama': d.get('agama', ''),
        'kewarganegaraan': d.get('kewarganegaraan', ''),
        'status_kahwin': d.get('status_kahwin', ''),
        'status_pendidikan': d.get('status_pendidikan', ''),
        'pekerjaan': d.get('pekerjaan', ''),
        'diagnosa': d.get('diagnosa', ''),
        'tarikh_masuk': d.get('tarikh_masuk', ''),
        'tarikh_keluar': d.get('tarikh_keluar', ''),
        'kategori_kes': d.get('kategori_kes', ''),

        # 2. Pendapatan & Perbelanjaan
        'pendapatan_isi_rumah': d.get('pendapatan_isi_rumah', '0'),
        'sumbangan_keluarga': d.get('sumbangan_keluarga', '0'),
        'bantuan_tetap': d.get('bantuan_tetap', '0'),
        'bantuan_jkm_penjaga': d.get('bantuan_jkm_penjaga', '0'),
        'bantuan_maidam': d.get('bantuan_maidam', '0'),
        'jumlah_pendapatan': d.get('jumlah_pendapatan', '0'),
        'belanja_makanan': d.get('belanja_makanan', '0'),
        'belanja_sewa': d.get('belanja_sewa', '0'),
        'belanja_utiliti': d.get('belanja_utiliti', '0'),
        'belanja_pengangkutan': d.get('belanja_pengangkutan', '0'),
        'belanja_pendidikan': d.get('belanja_pendidikan', '0'),
        'belanja_perubatan': d.get('belanja_perubatan', '0'),
        'belanja_pakaian': d.get('belanja_pakaian', '0'),
        'belanja_penjagaan_anak': d.get('belanja_penjagaan_anak', '0'),
        'belanja_pinjaman': d.get('belanja_pinjaman', '0'),
        'belanja_lain': d.get('belanja_lain', '0'),
        'belanja_lain_text': d.get('belanja_lain_text', ''),
        'jumlah_perbelanjaan': d.get('jumlah_perbelanjaan', '0'),
        'caruman_kwsp': d.get('caruman_kwsp', '0'),
        'caruman_perkeso': d.get('caruman_perkeso', '0'),
        'jumlah_caruman': d.get('jumlah_caruman', '0'),
        'tabungan_insuran': d.get('tabungan_insuran', '0'),
        'tabungan_asm': d.get('tabungan_asm', '0'),
        'tabungan_lain': d.get('tabungan_lain', '0'),
        'jumlah_tabungan': d.get('jumlah_tabungan', '0'),
        'harta_simpanan': d.get('harta_simpanan', '0'),
        'harta_rumah': d.get('harta_rumah', '0'),
        'harta_tanah': d.get('harta_tanah', '0'),
        'harta_kenderaan': d.get('harta_kenderaan', '0'),
        'harta_kenderaan_jenis': d.get('harta_kenderaan_jenis', ''),
        'harta_lain': d.get('harta_lain', '0'),
        'harta_lain_text': d.get('harta_lain_text', ''),
        'jumlah_harta': d.get('jumlah_harta', '0'),
        'baki_jumlah_pendapatan': d.get('baki_jumlah_pendapatan', '0'),
        'baki_jumlah_perbelanjaan': d.get('baki_jumlah_perbelanjaan', '0'),
        'baki_bersih': d.get('baki_bersih', '0'),

        # 3. Tempat Tinggal
        'jenis_kediaman': d.get('jenis_kediaman', ''),
        'jenis_kediaman_lain': d.get('jenis_kediaman_lain', ''),
        'status_kediaman': d.get('status_kediaman', ''),
        'utiliti': ', '.join(d.getlist('utiliti')) if isinstance(d.get('utiliti'), str) else ', '.join(d.getlist('utiliti')) if hasattr(d, 'getlist') else d.get('utiliti', ''),
        'kemudahan_lain': d.get('kemudahan_lain', ''),

        # 4. Keluarga Serumah — JSON string
        'keluarga_serumah': d.get('keluarga_serumah_json', '[]'),
        'keluarga_berasingan': d.get('keluarga_berasingan_json', '[]'),

        # Tanggungan text
        'catatan_tanggungan': d.get('catatan_tanggungan', ''),

        # 5. Sejarah Perubatan
        'sejarah_perubatan': d.get('sejarah_perubatan', ''),

        # 6. Aspek Sosial
        'tingkah_laku': d.get('tingkah_laku', ''),
        'penyakit_mental_keluarga': d.get('penyakit_mental_keluarga', ''),
        'sokongan_keluarga': d.get('sokongan_keluarga', ''),
        'pengalaman_kerja': d.get('pengalaman_kerja', ''),
        'pendidikan_pesakit': d.get('pendidikan_pesakit', ''),
        'perkahwinan_pesakit': d.get('perkahwinan_pesakit', ''),
        'rekod_jenayah': d.get('rekod_jenayah', ''),

        # 7. Ulasan PKSP
        'kesan_fizikal': d.get('kesan_fizikal', ''),
        'kesan_ekonomi': d.get('kesan_ekonomi', ''),
        'kesan_emosi': d.get('kesan_emosi', ''),
        'kesan_mental': d.get('kesan_mental', ''),
        'kesan_perhubungan': d.get('kesan_perhubungan', ''),
        'bantuan_kewangan_jenis': d.get('bantuan_kewangan_jenis', ''),
        'bantuan_kewangan_kos': d.get('bantuan_kewangan_kos', ''),
        'bantuan_kewangan_bayar': d.get('bantuan_kewangan_bayar', ''),
        'bantuan_kewangan_dipohon': d.get('bantuan_kewangan_dipohon', ''),
        'bantuan_kewangan_jangka': d.get('bantuan_kewangan_jangka', ''),
        'kesan_waris_laporan': d.get('kesan_waris_laporan', ''),
        'kesan_waris_agensi': d.get('kesan_waris_agensi', ''),
        'kesan_waris_siaram': d.get('kesan_waris_siaram', ''),
        'kesan_waris_jangka': d.get('kesan_waris_jangka', ''),
        'penempatan_institusi': d.get('penempatan_institusi', ''),
        'penempatan_bayaran': d.get('penempatan_bayaran', ''),
        'penempatan_alat_sokongan': d.get('penempatan_alat_sokongan', ''),
        'penempatan_pengangkutan': d.get('penempatan_pengangkutan', ''),
        'penempatan_jangka': d.get('penempatan_jangka', ''),
        'terapi_jenis': d.get('terapi_jenis', ''),
        'terapi_jangka': d.get('terapi_jangka', ''),
        'terapi_rujukan': d.get('terapi_rujukan', ''),

        # 8. Syor
        'syor': d.get('syor', ''),

        # 9. Tandatangan
        'nama_pk': d.get('nama_pk', ''),
        'jawatan_pk': d.get('jawatan_pk', 'Pegawai Kerja Sosial Perubatan'),
        'tarikh_laporan': d.get('tarikh_laporan', datetime.now().strftime('%d/%m/%Y')),
        'nama_ketua': d.get('nama_ketua', ''),
        'cop_ketua': d.get('cop_ketua', ''),
        'tarikh_ketua': d.get('tarikh_ketua', datetime.now().strftime('%d/%m/%Y')),
    }


# ─────────────────────────────────────────────────
# Jana .docx
# ─────────────────────────────────────────────────
def generate_docx(form_data):
    data = prepare_data(form_data)
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(3.5)      # lebih ruang untuk header gambar
        section.bottom_margin = Cm(3.0)   # lebih ruang untuk footer gambar
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Default style — Arial 11pt, line spacing 1.15, no space after ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # Also set heading styles
    for level in [1, 2, 3]:
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Arial'
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.paragraph_format.space_before = Pt(6)
        hs.paragraph_format.space_after = Pt(3)
        hs.paragraph_format.line_spacing = 1.15

    # ── Laporan BPS: TANPA letterhead (ruang header kosong utk cop/tandatangan) ──

    # ── Helper functions ──
    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        return h

    def add_field(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(f'{label}\t: ')
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        r2 = p.add_run(str(value) if value else '___________________')
        r2.font.name = 'Arial'
        r2.font.size = Pt(11)

    def add_section_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'

    def add_sub_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'

    # ═══════════════════════════════════
    # Tajuk
    # ═══════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SULIT')
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('LAPORAN PENILAIAN BIOPSIKOSOSIAL\n'
                     'PEGAWAI KERJA SOSIAL PERUBATAN (PKSP)\n'
                     'KEMENTERIAN KESIHATAN MALAYSIA')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()  # spacing

    # ═══════════════════════════════════
    # 1. MAKLUMAT PESAKIT
    # ═══════════════════════════════════
    add_heading('1. MAKLUMAT PESAKIT')
    add_field('No. Rujukan Fail', data['rujukan_fail'])
    add_field('Tarikh Pendaftaran', data['tarikh_daftar'])
    add_field('Nama', data['nama'])
    add_field('No. Kad Pengenalan', data['ic'])
    add_field('Alamat', data['alamat'])
    add_field('No. Telefon', data['telefon'])
    add_field('Tarikh Lahir/Umur', f"{data['tarikh_lahir']}/{data['umur']}" if data['umur'] else data['tarikh_lahir'])
    add_field('Tempat Lahir', data['tempat_lahir'])
    add_field('Jantina', data['jantina'])
    add_field('Bangsa', data['bangsa'])
    add_field('Agama', data['agama'])
    add_field('Kewarganegaraan', data['kewarganegaraan'])
    add_field('Status Perkahwinan', data['status_kahwin'])
    add_field('Status Pendidikan', data['status_pendidikan'])
    add_field('Pekerjaan', data['pekerjaan'])
    add_field('Diagnosa', data['diagnosa'])
    add_field('Tarikh Masuk Hospital', data['tarikh_masuk'])
    add_field('Tarikh Keluar Hospital', data['tarikh_keluar'])

    # ═══════════════════════════════════
    # 2. PENDAPATAN & PERBELANJAAN
    # ═══════════════════════════════════
    add_heading('2. PENDAPATAN DAN PERBELANJAAN DIRI/KELUARGA')
    add_sub_heading('2.1 Pendapatan Sebulan')
    add_field('i) Pendapatan Isi Rumah', f"RM {data['pendapatan_isi_rumah']}")
    add_field('ii) Sumbangan Ahli Keluarga', f"RM {data['sumbangan_keluarga']}")
    add_field('iii) Bantuan Tetap', f"RM {data['bantuan_tetap']}")
    add_field('iv) JKM-Penjaga', f"RM {data['bantuan_jkm_penjaga']}")
    add_field('MAIDAM', f"RM {data['bantuan_maidam']}")
    add_field('JUMLAH PENDAPATAN', f"RM {data['jumlah_pendapatan']}")

    add_sub_heading('2.2 Perbelanjaan Sebulan (Anggaran)')
    add_field('i) Makanan', f"RM {data['belanja_makanan']}")
    add_field('ii) Ansuran/Sewa Rumah', f"RM {data['belanja_sewa']}")
    add_field('iii) Utiliti', f"RM {data['belanja_utiliti']}")
    add_field('iv) Pengangkutan', f"RM {data['belanja_pengangkutan']}")
    add_field('v) Pendidikan', f"RM {data['belanja_pendidikan']}")
    add_field('vi) Kos Perubatan', f"RM {data['belanja_perubatan']}")
    add_field('vii) Pakaian', f"RM {data['belanja_pakaian']}")
    add_field('viii) Penjagaan Anak', f"RM {data['belanja_penjagaan_anak']}")
    add_field('ix) Pinjaman Peribadi', f"RM {data['belanja_pinjaman']}")
    add_field(f"x) Lain-lain ({data['belanja_lain_text']})" if data['belanja_lain_text'] else 'x) Lain-lain', f"RM {data['belanja_lain']}")
    add_field('JUMLAH PERBELANJAAN', f"RM {data['jumlah_perbelanjaan']}")

    add_sub_heading('2.3 Caruman Bulanan')
    add_field('i) KWSP', f"RM {data['caruman_kwsp']}")
    add_field('ii) PERKESO', f"RM {data['caruman_perkeso']}")
    add_field('JUMLAH CARUMAN', f"RM {data['jumlah_caruman']}")

    add_sub_heading('2.4 Tabungan Bulanan')
    add_field('i) Insuran', f"RM {data['tabungan_insuran']}")
    add_field('ii) Koperasi/ASN/ASB/Tabung Haji', f"RM {data['tabungan_asm']}")
    add_field('iii) Lain-Lain', f"RM {data['tabungan_lain']}")
    add_field('JUMLAH TABUNGAN', f"RM {data['jumlah_tabungan']}")

    add_sub_heading('2.5 Harta (Nilai)')
    add_field('i) Wang Simpanan/Saham', f"RM {data['harta_simpanan']}")
    add_field('ii) Rumah', f"RM {data['harta_rumah']}")
    add_field('iii) Tanah', f"RM {data['harta_tanah']}")
    jenis_kereta = f" ({data['harta_kenderaan_jenis']})" if data['harta_kenderaan_jenis'] else ''
    add_field(f'iv) Kenderaan{jenis_kereta}', f"RM {data['harta_kenderaan']}")
    lain_harta = f" ({data['harta_lain_text']})" if data['harta_lain_text'] else ''
    add_field(f'v) Lain-lain{lain_harta}', f"RM {data['harta_lain']}")
    add_field('JUMLAH HARTA', f"RM {data['jumlah_harta']}")

    add_sub_heading('2.6 Baki Pendapatan Bulanan')
    add_field('i) Jumlah Pendapatan', f"RM {data['baki_jumlah_pendapatan']}")
    add_field('ii) Jumlah Perbelanjaan', f"RM {data['baki_jumlah_perbelanjaan']}")
    add_field('BAKI BERSIH', f"RM {data['baki_bersih']}")

    # ═══════════════════════════════════
    # 3. KEADAAN TEMPAT TINGGAL
    # ═══════════════════════════════════
    add_heading('3. KEADAAN TEMPAT TINGGAL')
    add_field('Jenis Kediaman', data['jenis_kediaman'] + (f" ({data['jenis_kediaman_lain']})" if data['jenis_kediaman'] == 'Lain-lain' and data['jenis_kediaman_lain'] else ''))
    add_field('Status Kediaman', data['status_kediaman'])
    add_field('Utiliti', data['utiliti'] if data['utiliti'] else 'Tiada')
    if data['kemudahan_lain']:
        add_field('Kemudahan Lain', data['kemudahan_lain'])

    # ═══════════════════════════════════
    # 4. BUTIR KELUARGA
    # ═══════════════════════════════════
    add_heading('4. BUTIR-BUTIR KELUARGA')
    add_sub_heading('4.1 Keluarga Yang Tinggal Serumah')

    try:
        keluarga_serumah = json.loads(data['keluarga_serumah'])
    except (json.JSONDecodeError, TypeError):
        keluarga_serumah = []

    if keluarga_serumah:
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        headers = ['Bil', 'Nama', 'Jantina/Umur', 'Status', 'Pertalian', 'Pekerjaan/Sekolah', 'Pendapatan (RM)']
        for i, h_text in enumerate(headers):
            hdr[i].text = h_text
            for para in hdr[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.bold = True
        for idx, ahli in enumerate(keluarga_serumah, 1):
            row = table.add_row().cells
            values = [
                str(idx),
                ahli.get('nama', ''),
                f"{ahli.get('jantina', '')}/{ahli.get('umur', '')}",
                ahli.get('status', ''),
                ahli.get('pertalian', ''),
                ahli.get('pekerjaan', ''),
                ahli.get('pendapatan', ''),
            ]
            for i, val in enumerate(values):
                row[i].text = val
                for para in row[i].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
    else:
        doc.add_paragraph('(Tiada maklumat)')

    add_sub_heading('4.2 Keluarga Yang Tinggal Berasingan')
    try:
        keluarga_berasingan = json.loads(data['keluarga_berasingan'])
    except (json.JSONDecodeError, TypeError):
        keluarga_berasingan = []

    if keluarga_berasingan:
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h_text in enumerate(headers):
            hdr[i].text = h_text
            for para in hdr[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.bold = True
        for idx, ahli in enumerate(keluarga_berasingan, 1):
            row = table.add_row().cells
            values = [
                str(idx),
                ahli.get('nama', ''),
                f"{ahli.get('jantina', '')}/{ahli.get('umur', '')}",
                ahli.get('status', ''),
                ahli.get('pertalian', ''),
                ahli.get('pekerjaan', ''),
                ahli.get('pendapatan', ''),
            ]
            for i, val in enumerate(values):
                row[i].text = val
                for para in row[i].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
    else:
        doc.add_paragraph('(Tiada maklumat)')

    if data['catatan_tanggungan']:
        doc.add_paragraph('')
        add_field('Catatan Tanggungan', data['catatan_tanggungan'])

    # ═══════════════════════════════════
    # 5. SEJARAH PERUBATAN
    # ═══════════════════════════════════
    add_heading('5. SEJARAH PERUBATAN PESAKIT')
    doc.add_paragraph('(Kecacatan, masalah kesihatan, masalah kesihatan mental, penyalahgunaan dadah/alkohol)')
    if data['sejarah_perubatan']:
        doc.add_paragraph(data['sejarah_perubatan'])
    else:
        doc.add_paragraph('_________________________')

    # ═══════════════════════════════════
    # 6. ASPEK SOSIAL
    # ═══════════════════════════════════
    add_heading('6. MAKLUMAT TAMBAHAN ASPEK SOSIAL')

    sosial_fields = [
        ('a) Tingkah laku di rumah dan dalam komuniti', data['tingkah_laku']),
        ('b) Ahli keluarga yang menghidap penyakit mental', data['penyakit_mental_keluarga']),
        ('c) Sistem Sokongan Keluarga\n   (Masalah penjagaan, hubungan, penerimaan, ekspektasi keluarga)', data['sokongan_keluarga']),
        ('d) Pengalaman bekerja', data['pengalaman_kerja']),
        ('e) Pendidikan', data['pendidikan_pesakit']),
        ('f) Perkahwinan', data['perkahwinan_pesakit']),
        ('g) Rekod jenayah', data['rekod_jenayah']),
    ]
    for label, value in sosial_fields:
        add_sub_heading(label)
        if value:
            doc.add_paragraph(value)
        else:
            doc.add_paragraph('_________________________')

    # ═══════════════════════════════════
    # 7. ULASAN
    # ═══════════════════════════════════
    add_heading('7. ULASAN PEGAWAI KERJA SOSIAL PERUBATAN')

    add_sub_heading('7.1 Kesan Penyakit Terhadap Kefungsian Sosial Pesakit/Keluarga')
    kesan_fields = [
        ('a) Fizikal', data['kesan_fizikal']),
        ('b) Ekonomi', data['kesan_ekonomi']),
        ('c) Emosi', data['kesan_emosi']),
        ('d) Mental', data['kesan_mental']),
        ('e) Perhubungan Sosial', data['kesan_perhubungan']),
    ]
    for label, value in kesan_fields:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(11)
        if value:
            doc.add_paragraph(value)
        else:
            doc.add_paragraph('_________________________')

    add_sub_heading('7.2 Bantuan Yang Disyorkan')
    add_sub_heading('a) Bantuan Praktik')

    # i) Bantuan Kewangan
    p = doc.add_paragraph()
    run = p.add_run('i) Bantuan Kewangan (Peralatan Perubatan/Ubat/Rawatan)')
    run.bold = True
    run.font.size = Pt(10)
    add_field('Jenis peralatan/ubat/rawatan', data['bantuan_kewangan_jenis'])
    add_field('Kos', f"RM {data['bantuan_kewangan_kos']}" if data['bantuan_kewangan_kos'] else '')
    add_field('Jumlah dibayar pesakit', data['bantuan_kewangan_bayar'])
    add_field('Jumlah bantuan dipohon', f"RM {data['bantuan_kewangan_dipohon']}" if data['bantuan_kewangan_dipohon'] else '')
    add_field('Jangka masa diperlukan', data['bantuan_kewangan_jangka'])

    # ii) Kesan Waris
    p = doc.add_paragraph()
    run = p.add_run('ii) Kesan Waris')
    run.bold = True
    run.font.size = Pt(10)
    add_field('Laporan Polis', data['kesan_waris_laporan'])
    add_field('Agensi Rujukan', data['kesan_waris_agensi'])
    add_field('Siaran Akhbar', data['kesan_waris_siaram'])
    add_field('Jangka masa', data['kesan_waris_jangka'])

    # iii) Penempatan Institusi
    p = doc.add_paragraph()
    run = p.add_run('iii) Penempatan Institusi')
    run.bold = True
    run.font.size = Pt(10)
    add_field('Cadangan Institusi', data['penempatan_institusi'])
    add_field('Bayaran', data['penempatan_bayaran'])
    add_field('Alat Sokongan/Rawatan', data['penempatan_alat_sokongan'])
    add_field('Pengangkutan/Pengiring', data['penempatan_pengangkutan'])
    add_field('Jangka masa', data['penempatan_jangka'])

    # b) Terapi Sokongan
    add_sub_heading('b) Terapi Sokongan (Khidmat Perundingan/Sokongan Emosi/Intervensi Krisis)')
    add_field('Jenis Terapi', data['terapi_jenis'])
    add_field('Jangka masa', data['terapi_jangka'])
    add_field('Rujukan Agensi', data['terapi_rujukan'])

    # ═══════════════════════════════════
    # 8. SYOR
    # ═══════════════════════════════════
    add_heading('8. SYOR/PELAN TINDAKAN PEGAWAI KERJA SOSIAL PERUBATAN')
    if data['syor']:
        doc.add_paragraph(data['syor'])
    else:
        doc.add_paragraph('_________________________')

    # ═══════════════════════════════════
    # 9. TANDATANGAN
    # ═══════════════════════════════════
    doc.add_paragraph()
    add_heading('TANDATANGAN')
    add_field('Tandatangan', '')
    add_field('Nama & Cop Pegawai', data['nama_pk'])
    add_field('Jawatan', data['jawatan_pk'])
    add_field('Tarikh', data['tarikh_laporan'])

    doc.add_paragraph()
    add_heading('ULASAN DAN PENGESAHAN KETUA JABATAN / PEGAWAI KANAN')
    p = doc.add_paragraph()
    run = p.add_run('Disokong')
    run.bold = True
    add_field('Tandatangan', '')
    add_field('Nama & Cop Pegawai', data['nama_ketua'] if data['nama_ketua'] else '___________________')
    add_field('Tarikh', data['tarikh_ketua'])

    # Save
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    nama_pendek = data['nama'].replace(' ', '_')[:30] if data['nama'] else 'unknown'

    # Simpan ke Dropbox: Laporan BPS <Tahun>/<Nama Pesakit>/
    nama_folder = ' '.join(data['nama'].strip().split()) if data['nama'] else 'Unknown'
    nama_folder = nama_folder.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
    patient_dir = os.path.join(DROPBOX_BPS, nama_folder)
    os.makedirs(patient_dir, exist_ok=True)

    fname = f"LAPORAN BPS {data['nama'].upper() if data['nama'] else 'UNKNOWN'} {timestamp}.docx"
    fpath = os.path.join(patient_dir, fname)
    doc.save(fpath)

    return send_file(fpath, as_attachment=True, download_name=fname)


# ─────────────────────────────────────────────────
# Helper: bina HTML laporan (untuk PDF / print-friendly)
# ─────────────────────────────────────────────────
def build_report_html(data):
    try:
        keluarga_serumah = json.loads(data['keluarga_serumah'])
    except:
        keluarga_serumah = []
    try:
        keluarga_berasingan = json.loads(data['keluarga_berasingan'])
    except:
        keluarga_berasingan = []

    def esc(s):
        if s is None: return ''
        return str(s).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')

    def field(label, value, unit=''):
        v = esc(value) if value else '<span class="kosong">___________________</span>'
        prefix = f'{unit} ' if unit else ''
        return f'<div class="field"><b>{esc(label)}</b> : {prefix}{v}</div>'

    def naratif(value):
        return esc(value) if value else '<span class="kosong">_________________________</span>'

    def table_keluarga(ahli_list):
        if not ahli_list:
            return '<p class="kosong">(Tiada maklumat)</p>'
        rows = '<tr><th>Bil</th><th>Nama</th><th>Jantina/Umur</th><th>Status</th><th>Pertalian</th><th>Pekerjaan/Sekolah</th><th>Pendapatan (RM)</th></tr>'
        for idx, a in enumerate(ahli_list, 1):
            rows += f'<tr><td>{idx}</td><td>{esc(a.get("nama",""))}</td><td>{esc(a.get("jantina",""))}/{esc(a.get("umur",""))}</td><td>{esc(a.get("status",""))}</td><td>{esc(a.get("pertalian",""))}</td><td>{esc(a.get("pekerjaan",""))}</td><td>{esc(a.get("pendapatan",""))}</td></tr>'
        return f'<table>{rows}</table>'

    jk = data['jenis_kediaman']
    if jk == 'Lain-lain' and data['jenis_kediaman_lain']:
        jk += f' ({data["jenis_kediaman_lain"]})'

    harta_kend = f'iv) Kenderaan{" (" + data["harta_kenderaan_jenis"] + ")" if data["harta_kenderaan_jenis"] else ""}'
    harta_lain = f'v) Lain-lain{" (" + data["harta_lain_text"] + ")" if data["harta_lain_text"] else ""}'
    bl_lain = f'x) Lain-lain{" (" + data["belanja_lain_text"] + ")" if data["belanja_lain_text"] else ""}'

    # Kategori kes
    kategori = data.get('kategori_kes', '')

    return f'''<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
@page {{ size: A4; margin: 2.5cm; }}
body {{ font-family: Arial, sans-serif; font-size: 11pt; line-height: 1.5; }}
.sulit {{ text-align: center; font-weight: bold; }}
.tajuk {{ text-align: center; font-weight: bold; font-size: 13pt; margin-bottom: 20px; }}
h1 {{ font-size: 12pt; margin-top: 16px; margin-bottom: 6px; border-bottom: 1px solid #333; }}
h2 {{ font-size: 11pt; margin-top: 12px; margin-bottom: 4px; font-style: italic; }}
.field {{ margin: 2px 0; }}
.field b {{ display: inline-block; min-width: 230px; }}
table {{ width: 100%; border-collapse: collapse; margin: 8px 0; font-size: 9pt; }}
th, td {{ border: 1px solid #555; padding: 3px 5px; text-align: left; }}
th {{ background: #e0e0e0; font-weight: bold; }}
.naratif {{ margin: 4px 0 10px 0; padding-left: 10px; white-space: pre-wrap; }}
.kosong {{ color: #999; font-style: italic; }}
@media print {{ .no-print {{ display: none; }} }}
</style></head><body>

<p class="sulit">SULIT</p>
<p class="tajuk">LAPORAN PENILAIAN BIOPSIKOSOSIAL<br>
PEGAWAI KERJA SOSIAL PERUBATAN (PKSP)<br>
KEMENTERIAN KESIHATAN MALAYSIA</p>

{f'<p><b>Kategori Kes :</b> {esc(kategori)}</p>' if kategori else ''}

<h1>1. MAKLUMAT PESAKIT</h1>
{field('No. Rujukan Fail', data['rujukan_fail'])}
{field('Tarikh Pendaftaran', data['tarikh_daftar'])}
{field('Nama', data['nama'])}
{field('No. Kad Pengenalan', data['ic'])}
{field('Alamat', data['alamat'])}
{field('No. Telefon', data['telefon'])}
{field('Tarikh Lahir/Umur', f"{data['tarikh_lahir']}/{data['umur']}" if data['umur'] else data['tarikh_lahir'])}
{field('Tempat Lahir', data['tempat_lahir'])}
{field('Jantina', data['jantina'])}
{field('Bangsa', data['bangsa'])}
{field('Agama', data['agama'])}
{field('Kewarganegaraan', data['kewarganegaraan'])}
{field('Status Perkahwinan', data['status_kahwin'])}
{field('Status Pendidikan', data['status_pendidikan'])}
{field('Pekerjaan', data['pekerjaan'])}
{field('Diagnosa', data['diagnosa'])}
{field('Tarikh Masuk Hospital', data['tarikh_masuk'])}
{field('Tarikh Keluar Hospital', data['tarikh_keluar'])}

<h1>2. PENDAPATAN DAN PERBELANJAAN DIRI/KELUARGA</h1>
<h2>2.1 Pendapatan Sebulan</h2>
{field('i) Pendapatan Isi Rumah', data['pendapatan_isi_rumah'], 'RM')}
{field('ii) Sumbangan Ahli Keluarga', data['sumbangan_keluarga'], 'RM')}
{field('iii) Bantuan Tetap', data['bantuan_tetap'], 'RM')}
{field('iv) JKM-Penjaga', data['bantuan_jkm_penjaga'], 'RM')}
{field('MAIDAM', data['bantuan_maidam'], 'RM')}
{field('JUMLAH PENDAPATAN', data['jumlah_pendapatan'], 'RM')}

<h2>2.2 Perbelanjaan Sebulan (Anggaran)</h2>
{field('i) Makanan', data['belanja_makanan'], 'RM')}
{field('ii) Ansuran/Sewa Rumah', data['belanja_sewa'], 'RM')}
{field('iii) Utiliti', data['belanja_utiliti'], 'RM')}
{field('iv) Pengangkutan', data['belanja_pengangkutan'], 'RM')}
{field('v) Pendidikan', data['belanja_pendidikan'], 'RM')}
{field('vi) Kos Perubatan', data['belanja_perubatan'], 'RM')}
{field('vii) Pakaian', data['belanja_pakaian'], 'RM')}
{field('viii) Penjagaan Anak', data['belanja_penjagaan_anak'], 'RM')}
{field('ix) Pinjaman Peribadi', data['belanja_pinjaman'], 'RM')}
{field(bl_lain, data['belanja_lain'], 'RM')}
{field('JUMLAH PERBELANJAAN', data['jumlah_perbelanjaan'], 'RM')}

<h2>2.3 Caruman Bulanan</h2>
{field('i) KWSP', data['caruman_kwsp'], 'RM')}
{field('ii) PERKESO', data['caruman_perkeso'], 'RM')}
{field('JUMLAH CARUMAN', data['jumlah_caruman'], 'RM')}

<h2>2.4 Tabungan Bulanan</h2>
{field('i) Insuran', data['tabungan_insuran'], 'RM')}
{field('ii) Koperasi/ASN/ASB/Tabung Haji', data['tabungan_asm'], 'RM')}
{field('iii) Lain-Lain', data['tabungan_lain'], 'RM')}
{field('JUMLAH TABUNGAN', data['jumlah_tabungan'], 'RM')}

<h2>2.5 Harta (Nilai)</h2>
{field('i) Wang Simpanan/Saham', data['harta_simpanan'], 'RM')}
{field('ii) Rumah', data['harta_rumah'], 'RM')}
{field('iii) Tanah', data['harta_tanah'], 'RM')}
{field(harta_kend, data['harta_kenderaan'], 'RM')}
{field(harta_lain, data['harta_lain'], 'RM')}
{field('JUMLAH HARTA', data['jumlah_harta'], 'RM')}

<h2>2.6 Baki Pendapatan Bulanan</h2>
{field('i) Jumlah Pendapatan', data['baki_jumlah_pendapatan'], 'RM')}
{field('ii) Jumlah Perbelanjaan', data['baki_jumlah_perbelanjaan'], 'RM')}
{field('BAKI BERSIH', data['baki_bersih'], 'RM')}

<h1>3. KEADAAN TEMPAT TINGGAL</h1>
{field('Jenis Kediaman', jk)}
{field('Status Kediaman', data['status_kediaman'])}
{field('Utiliti', data['utiliti'] or 'Tiada')}
{f'<div class="field"><b>Kemudahan Lain</b> : {esc(data["kemudahan_lain"])}</div>' if data['kemudahan_lain'] else ''}

<h1>4. BUTIR-BUTIR KELUARGA</h1>
<h2>4.1 Keluarga Yang Tinggal Serumah</h2>
{table_keluarga(keluarga_serumah)}
<h2>4.2 Keluarga Yang Tinggal Berasingan</h2>
{table_keluarga(keluarga_berasingan)}
{f'<div class="field"><b>Catatan Tanggungan</b> : {esc(data["catatan_tanggungan"])}</div>' if data['catatan_tanggungan'] else ''}

<h1>5. SEJARAH PERUBATAN PESAKIT</h1>
<p><i>(Kecacatan, masalah kesihatan, masalah kesihatan mental, penyalahgunaan dadah/alkohol)</i></p>
<div class="naratif">{naratif(data['sejarah_perubatan'])}</div>

<h1>6. MAKLUMAT TAMBAHAN ASPEK SOSIAL</h1>
<h2>a) Tingkah laku di rumah dan dalam komuniti</h2><div class="naratif">{naratif(data['tingkah_laku'])}</div>
<h2>b) Ahli keluarga yang menghidap penyakit mental</h2><div class="naratif">{naratif(data['penyakit_mental_keluarga'])}</div>
<h2>c) Sistem Sokongan Keluarga</h2><div class="naratif">{naratif(data['sokongan_keluarga'])}</div>
<h2>d) Pengalaman bekerja</h2><div class="naratif">{naratif(data['pengalaman_kerja'])}</div>
<h2>e) Pendidikan</h2><div class="naratif">{naratif(data['pendidikan_pesakit'])}</div>
<h2>f) Perkahwinan</h2><div class="naratif">{naratif(data['perkahwinan_pesakit'])}</div>
<h2>g) Rekod jenayah</h2><div class="naratif">{naratif(data['rekod_jenayah'])}</div>

<h1>7. ULASAN PEGAWAI KERJA SOSIAL PERUBATAN</h1>
<h2>7.1 Kesan Penyakit Terhadap Kefungsian Sosial Pesakit/Keluarga</h2>
<h2>a) Fizikal</h2><div class="naratif">{naratif(data['kesan_fizikal'])}</div>
<h2>b) Ekonomi</h2><div class="naratif">{naratif(data['kesan_ekonomi'])}</div>
<h2>c) Emosi</h2><div class="naratif">{naratif(data['kesan_emosi'])}</div>
<h2>d) Mental</h2><div class="naratif">{naratif(data['kesan_mental'])}</div>
<h2>e) Perhubungan Sosial</h2><div class="naratif">{naratif(data['kesan_perhubungan'])}</div>

<h2>7.2 Bantuan Yang Disyorkan</h2>
<h2>a) Bantuan Praktik</h2>
<h2>i) Bantuan Kewangan (Peralatan Perubatan/Ubat/Rawatan)</h2>
{field('Jenis peralatan/ubat/rawatan', data['bantuan_kewangan_jenis'])}
{field('Kos', data['bantuan_kewangan_kos'], 'RM')}
{field('Jumlah dibayar pesakit', data['bantuan_kewangan_bayar'])}
{field('Jumlah bantuan dipohon', data['bantuan_kewangan_dipohon'], 'RM')}
{field('Jangka masa', data['bantuan_kewangan_jangka'])}

<h2>ii) Kesan Waris</h2>
{field('Laporan Polis', data['kesan_waris_laporan'])}
{field('Agensi Rujukan', data['kesan_waris_agensi'])}
{field('Siaran Akhbar', data['kesan_waris_siaram'])}
{field('Jangka masa', data['kesan_waris_jangka'])}

<h2>iii) Penempatan Institusi</h2>
{field('Cadangan Institusi', data['penempatan_institusi'])}
{field('Bayaran', data['penempatan_bayaran'])}
{field('Alat Sokongan/Rawatan', data['penempatan_alat_sokongan'])}
{field('Pengangkutan/Pengiring', data['penempatan_pengangkutan'])}
{field('Jangka masa', data['penempatan_jangka'])}

<h2>b) Terapi Sokongan (Khidmat Perundingan/Sokongan Emosi/Intervensi Krisis)</h2>
{field('Jenis Terapi', data['terapi_jenis'])}
{field('Jangka masa', data['terapi_jangka'])}
{field('Rujukan Agensi', data['terapi_rujukan'])}

<h1>8. SYOR/PELAN TINDAKAN PEGAWAI KERJA SOSIAL PERUBATAN</h1>
<div class="naratif">{naratif(data['syor'])}</div>

<h1>TANDATANGAN</h1>
{field('Tandatangan', '')}
{field('Nama & Cop Pegawai', data['nama_pk'])}
{field('Jawatan', data['jawatan_pk'])}
{field('Tarikh', data['tarikh_laporan'])}

<h1>ULASAN DAN PENGESAHAN KETUA JABATAN / PEGAWAI KANAN</h1>
<div class="field"><b>Disokong</b></div>
{field('Tandatangan', '')}
{field('Nama & Cop Pegawai', data['nama_ketua'] or '___________________')}
{field('Tarikh', data['tarikh_ketua'])}

<p class="no-print" style="margin-top:30px;text-align:center;color:#888;font-size:10pt;">
── Laporan ini dijana oleh BPS Report Generator. Guna Ctrl+P → Save as PDF untuk simpan. ──
</p>
</body></html>'''


# ─────────────────────────────────────────────────
# Jana PDF
# ─────────────────────────────────────────────────
def generate_pdf(form_data):
    data = prepare_data(form_data)
    html = build_report_html(data)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    # Simpan ke Dropbox: Laporan BPS <Tahun>/<Nama Pesakit>/
    nama_folder = ' '.join(data['nama'].strip().split()) if data['nama'] else 'Unknown'
    nama_folder = nama_folder.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
    patient_dir = os.path.join(DROPBOX_BPS, nama_folder)
    os.makedirs(patient_dir, exist_ok=True)

    fname = f"LAPORAN BPS {data['nama'].upper() if data['nama'] else 'UNKNOWN'} {timestamp}.pdf"
    fpath = os.path.join(patient_dir, fname)

    if HAVE_WEASYPRINT:
        HTML(string=html).write_pdf(fpath)
        return send_file(fpath, as_attachment=True, download_name=fname)
    else:
        # Fallback: serve as HTML — user can Ctrl+P → Save as PDF
        return render_template('report.html', html_content=html)


# ─────────────────────────────────────────────────
# API: Generate Syor secara pintar
# ─────────────────────────────────────────────────
@app.route('/generate_syor', methods=['POST'])
def generate_syor():
    data = prepare_data(request.form)
    user_text = request.form.get('syor', '').strip()
    
    parts = []
    
    # 1. Maklumat pesakit
    nama = data['nama']
    umur = data['umur']
    jantina = data['jantina']
    diagnosa = data['diagnosa']
    
    intro = f"Pesakit {'bernama ' + nama if nama else ''}"
    if umur:
        intro += f", {umur} tahun"
    if jantina:
        intro += f", {jantina.lower()}"
    if diagnosa:
        intro += f", didiagnos {diagnosa}"
    if nama or umur or jantina or diagnosa:
        intro += "."
        parts.append(intro)
    
    # 2. Kategori kes / bantuan
    kat = data['kategori_kes']
    if kat:
        parts.append(f"Kes ini dikategorikan sebagai {kat}.")
    
    # 3. Pendapatan & perbelanjaan
    pendapatan = data['jumlah_pendapatan']
    perbelanjaan = data['jumlah_perbelanjaan']
    baki = data['baki_bersih']
    
    if pendapatan and pendapatan != '0':
        parts.append(f"Pendapatan isi rumah sebanyak RM{pendapatan}/bulan dengan perbelanjaan RM{perbelanjaan}/bulan," + 
                    (f" menjadikan baki bersih sebanyak RM{baki}/bulan." if baki != '0' else " mengalami defisit."))
    
    # 4. Tempat tinggal
    kediaman = data['jenis_kediaman']
    status_rumah = data['status_kediaman']
    if kediaman:
        parts.append(f"Pesakit menetap di {kediaman.lower()}" + (f" secara {status_rumah.lower()}." if status_rumah else "."))
    
    # 5. Keluarga (kira bilangan dari JSON)
    try:
        serumah = json.loads(data['keluarga_serumah'])
        num_serumah = len(serumah)
    except:
        num_serumah = 0
    try:
        berasingan = json.loads(data['keluarga_berasingan'])
        num_berasingan = len(berasingan)
    except:
        num_berasingan = 0
    
    if num_serumah > 0:
        parts.append(f"Tinggal serumah dengan {num_serumah} orang ahli keluarga.")
    if num_berasingan > 0:
        parts.append(f"Terdapat {num_berasingan} orang ahli keluarga yang tinggal berasingan.")
    
    # 6. Sokongan keluarga & sosial
    sokong = data['sokongan_keluarga']
    if sokong and len(sokong) > 10:
        parts.append(f"Dari segi sokongan keluarga: {sokong.strip().rstrip('.')}.")
    
    tingkah = data['tingkah_laku']
    if tingkah and len(tingkah) > 10:
        parts.append(f"Tingkah laku pesakit dalam komuniti: {tingkah.strip().rstrip('.')}.")
    
    # 7. Sejarah perubatan
    sejarah = data['sejarah_perubatan']
    if sejarah and len(sejarah) > 10:
        parts.append(f"Sejarah perubatan: {sejarah.strip().rstrip('.')}.")
    
    # 8. Kesan penyakit
    for label, key in [
        ('Fizikal', 'kesan_fizikal'),
        ('Ekonomi', 'kesan_ekonomi'),
        ('Emosi', 'kesan_emosi'),
    ]:
        val = data.get(key, '')
        if val and len(val) > 10:
            parts.append(f"Kesan {label.lower()}: {val.strip().rstrip('.')}.")
    
    # 9. Bantuan disyorkan
    bantuan = []
    if data['bantuan_kewangan_jenis']:
        bantuan.append(f"bantuan kewangan ({data['bantuan_kewangan_jenis']})")
    if data['penempatan_institusi']:
        bantuan.append(f"penempatan di {data['penempatan_institusi']}")
    if data['kesan_waris_agensi']:
        bantuan.append(f"kesan waris melalui {data['kesan_waris_agensi']}")
    if data['terapi_jenis']:
        bantuan.append(f"terapi sokongan ({data['terapi_jenis']})")
    
    if bantuan:
        parts.append(f"Bantuan yang disyorkan: {'; '.join(bantuan)}.")
    
    # Gabung
    syor = ' '.join(parts)
    
    # If user already typed something, prepend it (if it's not already in the generated text)
    if user_text:
        # Check if user_text is already contained in syor
        if user_text.lower() not in syor.lower():
            syor = user_text.rstrip('.') + '. ' + syor
        else:
            syor = syor  # User text already part of it
    
    # Trim to roughly 100 words
    words = syor.split()
    if len(words) > 150:
        syor = ' '.join(words[:150]) + '.'
    
    # If nothing was generated, provide a default
    if not syor.strip() or len(parts) == 0:
        syor = "Pesakit memerlukan penilaian dan bantuan sosial lanjut berdasarkan maklumat yang telah dikumpul. Sila lengkapkan maklumat di bahagian 1-7 untuk menjana laporan yang lebih terperinci."
    
    return jsonify({'syor': syor})


# ─────────────────────────────────────────────────
# API: Generate Syor dengan AI (OpenRouter)
# ─────────────────────────────────────────────────
@app.route('/generate_syor_ai', methods=['POST'])
def generate_syor_ai():
    data = prepare_data(request.form)
    user_text = request.form.get('syor', '').strip()

    if not OPENROUTER_API_KEY:
        return jsonify({'error': 'Tiada sambungan AI. PC ini mungkin offline atau tiada API key. Guna butang "Generate Syor (Offline)" untuk template.'}), 503

    # Kumpul semua maklumat untuk prompt
    try:
        serumah = json.loads(data['keluarga_serumah'])
    except:
        serumah = []
    try:
        berasingan = json.loads(data['keluarga_berasingan'])
    except:
        berasingan = []

    # Bina gambaran pesakit
    keluarga_desc = "Tiada"
    if serumah:
        keluarga_desc = ", ".join([f"{a.get('nama','?')} ({a.get('pertalian','?')}, {a.get('umur','?')}thn)" for a in serumah])

    info = f"""
MAKLUMAT PESAKIT:
- Nama: {data['nama']}
- Umur: {data['umur']}, Jantina: {data['jantina']}
- Diagnosa: {data['diagnosa']}
- Kategori Kes: {data['kategori_kes']}
- Pekerjaan: {data['pekerjaan']}
- Status kahwin: {data['status_kahwin']}, Pendidikan: {data['status_pendidikan']}

KEWANGAN:
- Pendapatan isi rumah: RM{data['pendapatan_isi_rumah']}
- Sumbangan keluarga: RM{data['sumbangan_keluarga']}
- JKM-Penjaga: RM{data['bantuan_jkm_penjaga']}, MAIDAM: RM{data['bantuan_maidam']}
- Jumlah pendapatan: RM{data['jumlah_pendapatan']}
- Jumlah perbelanjaan: RM{data['jumlah_perbelanjaan']}
- Baki bersih: RM{data['baki_bersih']}

TEMPAT TINGGAL:
- {data['jenis_kediaman']} ({data['status_kediaman']})
- Utiliti: {data['utiliti']}

KELUARGA:
- Serumah: {keluarga_desc}
- Bilangan serumah: {len(serumah)}, berasingan: {len(berasingan)}

SEJARAH PERUBATAN:
- {data['sejarah_perubatan']}

ASPEK SOSIAL:
- Tingkah laku: {data['tingkah_laku']}
- Sokongan keluarga: {data['sokongan_keluarga']}
- Pengalaman kerja: {data['pengalaman_kerja']}
- Pendidikan: {data['pendidikan_pesakit']}
- Perkahwinan: {data['perkahwinan_pesakit']}

KESAN PENYAKIT:
- Fizikal: {data['kesan_fizikal']}
- Ekonomi: {data['kesan_ekonomi']}
- Emosi: {data['kesan_emosi']}
- Mental: {data['kesan_mental']}
- Perhubungan sosial: {data['kesan_perhubungan']}

BANTUAN DISYORKAN:
- Bantuan kewangan jenis: {data['bantuan_kewangan_jenis']}
- Penempatan: {data['penempatan_institusi']}
- Terapi: {data['terapi_jenis']}
- Rujukan: {data['terapi_rujukan']}
"""

    user_extra = ""
    if user_text:
        user_extra = f"\n\nSentimen/ayat tambahan daripada pegawai sosial yang MESTI dimasukkan: \"{user_text}\""

    prompt = f"""Anda adalah Pratyeksia Kerja Sosial Perubatan (PKSP) di bawah Kementerian Kesihatan Malaysia. Tulis satu LAPORAN SYOR/PELAN TINDAKAN ringkas (~100 patah perkataan) dalam Bahasa Melayu formal-profesional untuk laporan penilaian biopsikososial.

Gunakan maklumat berikut:{info}{user_extra}

Tulis syor yang:
1. Ringkas, padat, ~100 patah perkataan
2. Bahasa Melayu formal (gaya laporan rasmi KKM)
3. Mulakan dengan ringkasan situasi pesakit
4. Kemudian nyatakan keperluan & cadangan bantuan
5. Akhiri dengan pelan tindakan / rujukan
6. JANGAN cipta maklumat yang tidak ada dalam data di atas
7. Jangan gunakan bullet point — tulis dalam bentuk perenggan naratif
"""

    headers = {
        'Authorization': f'Bearer {OPENROUTER_API_KEY}',
        'Content-Type': 'application/json',
        'HTTP-Referer': 'http://localhost:5000',
        'X-Title': 'BPS Report Generator'
    }
    payload = {
        'model': AI_MODEL,
        'messages': [{'role': 'user', 'content': prompt}],
        'temperature': 0.4,
        # v4-flash buat reasoning panjang; max_tokens kecil (400) sebabkan
        # finish_reason=length & content=None -> crash .strip()
        'max_tokens': 2000,
        # Matikan reasoning mode supaya output terus dalam 'content'
        # (kalau reasoning aktif, content boleh jadi None)
        'reasoning': {'enabled': False},
    }

    try:
        resp = requests.post('https://openrouter.ai/api/v1/chat/completions',
                             headers=headers, json=payload, timeout=60)
        resp.raise_for_status()
        result = resp.json()
        msg = result['choices'][0]['message']
        # DeepSeek v4-flash letak output dalam 'reasoning' bila reasoning mode;
        # 'content' boleh jadi None. Guna fallback supaya tak crash .strip().
        syor_raw = msg.get('content') or msg.get('reasoning') or ''
        if not syor_raw:
            return jsonify({'error': 'AI tak return kandungan (content kosong). '
                                      'Cuba lagi atau guna "Generate Syor (Offline)".'}), 502
        syor = syor_raw.strip()
        return jsonify({'syor': syor})
    except Exception as e:
        return jsonify({'error': f'Gagal panggil AI: {str(e)}'}), 500


# ─────────────────────────────────────────────────
# Generator Dokumen Tabung Darul Iman (TDI)
# ─────────────────────────────────────────────────
def _bulan_malaysia(dt):
    bulan = ['', 'JANUARI', 'FEBRUARI', 'MAC', 'APRIL', 'MEI', 'JUN',
             'JULAI', 'OGOS', 'SEPTEMBER', 'OKTOBER', 'NOVEMBER', 'DISEMBER']
    return bulan[dt.month]

def _format_tarikh(dt):
    return f"{dt.day} {_bulan_malaysia(dt)} {dt.year}"


def generate_surat_iringan(data, output_dir):
    """Jana SURAT IRINGAN TDI — dengan letterhead"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    for section in doc.sections:
        section.top_margin = Cm(3.5)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header & Footer letterhead
    hdr_img = _resource_path(os.path.join('static', 'letterhead-header.png'))
    ftr_img = _resource_path(os.path.join('static', 'letterhead-footer.png'))
    for section in doc.sections:
        if os.path.exists(hdr_img):
            h = section.header
            h.is_linked_to_previous = False
            hp = h.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hp.add_run()
            run.add_picture(hdr_img, width=Cm(15.5))
        if os.path.exists(ftr_img):
            f = section.footer
            f.is_linked_to_previous = False
            fp = f.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fp.add_run()
            run.add_picture(ftr_img, width=Cm(15.5))

    # Rujukan & Tarikh (kanan)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"Ruj. Kami : {data['rujukan']}\n")
    p.add_run(f"Tarikh    : {data['tarikh_tdi']}")

    doc.add_paragraph()

    # Alamat
    doc.add_paragraph('Pengurus')
    doc.add_paragraph('Sekretariat Tabung Darul Iman')
    doc.add_paragraph('Lot Pt. 6054, Tingkat Bawah, Satu & Dua')
    doc.add_paragraph('Dataran Alamanda. Jalan Sultan Sulaiman')
    doc.add_paragraph('2000 Kuala Terengganu')
    doc.add_paragraph('TERENGGANU DARUL IMAN')

    doc.add_paragraph()
    doc.add_paragraph('Tuan / Puan,')
    doc.add_paragraph()

    # Tajuk
    p = doc.add_paragraph()
    run = p.add_run('PENGESAHAN RAWATAN TERHADAP PERMOHONAN TABUNG DARUL IMAN')
    run.bold = True

    doc.add_paragraph(f"NAMA             : {data['nama']}")
    doc.add_paragraph(f"NO. K/P          : {data['ic']}")

    doc.add_paragraph()
    doc.add_paragraph('Dengan segala hormatnya saya merujuk kepada perkara di atas.')
    doc.add_paragraph()
    doc.add_paragraph('\tBersama-sama ini dilampirkan dokumen untuk perhatian pihak tuan.')

    doc.add_paragraph('Borang TDI')
    doc.add_paragraph('Ulasan Sosioekonomi (email: sectdi@gmail.com)')
    doc.add_paragraph('Salinan Kad Pengenalan Pesakit')

    doc.add_paragraph()
    doc.add_paragraph('\tPerhatian dan pertimbangan tuan terhadap permohonan ini didahului dengan ucapan ribuan terima kasih.')
    doc.add_paragraph()
    doc.add_paragraph('Sekian terima kasih.')

    doc.add_paragraph('"MALAYSIA MADANI"')
    doc.add_paragraph('"BERKHIDMAT UNTUK NEGARA"')

    doc.add_paragraph()
    doc.add_paragraph('Saya yang menjalankan amanah')
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(f'({data["nama_pk"] if data["nama_pk"] else NAMA_PEGAWAI})')
    run.bold = True
    doc.add_paragraph('Pegawai Kerja Sosial Perubatan')
    doc.add_paragraph('b.p Pengarah')
    doc.add_paragraph(NAMA_HOSPITAL)

    fpath = os.path.join(output_dir, 'SURAT IRINGAN.docx')
    doc.save(fpath)
    return fpath


def generate_borang_ulasan(data, output_dir):
    """Jana BORANG ULASAN TDI"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Tajuk
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ULASAN SOSIOEKONOMI')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # A. Butir-butir pesakit
    p = doc.add_paragraph()
    run = p.add_run('A. BUTIR – BUTIR PESAKIT')
    run.bold = True

    # Jadual 2 kolom supaya kolon : selari
    def _add_label_value_row(table, label, value):
        row = table.add_row().cells
        lp = row[0].paragraphs[0]
        lp.paragraph_format.line_spacing = 1.15
        lp.paragraph_format.space_after = Pt(2)
        lp.paragraph_format.space_before = Pt(2)
        lr = lp.add_run(label)
        lr.font.name = 'Arial'
        lr.font.size = Pt(11)
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vp = row[1].paragraphs[0]
        vp.paragraph_format.line_spacing = 1.15
        vp.paragraph_format.space_after = Pt(2)
        vp.paragraph_format.space_before = Pt(2)
        vp.add_run(': ' + str(value)).font.name = 'Arial'
        vp.runs[0].font.size = Pt(11)

    tbl = doc.add_table(rows=0, cols=2)
    tbl.autofit = True
    tbl.columns[0].width = Cm(7.5)
    tbl.columns[1].width = Cm(10.0)
    # label align kiri dengan tab-stop supaya label sama panjang
    _add_label_value_row(tbl, '1.  Nama', data['nama'])
    _add_label_value_row(tbl, '2.  No. K/P', data['ic'])
    _add_label_value_row(tbl, '3.  Diagnosis', data['diagnosa'])
    _add_label_value_row(tbl, '4.  Rawatan/ Peralatan', data['tdi_rawatan'])
    _add_label_value_row(tbl, '5.  Kos Rawatan / Peralatan (RM)', data['tdi_kos'])
    _add_label_value_row(tbl, '6.  Jumlah kos yang mampu ditanggung oleh pemohon/waris (RM)', data['tdi_mampu'])

    doc.add_paragraph()

    # B. Ulasan
    p = doc.add_paragraph()
    run = p.add_run('B.\tULASAN PEGAWAI KERJA SOSIAL PERUBATAN')
    run.bold = True

    ulasan = data['tdi_ulasan'] if data['tdi_ulasan'] else data['syor']
    doc.add_paragraph(ulasan if ulasan else '_________________________')

    doc.add_paragraph()

    # C. Syor
    p = doc.add_paragraph()
    run = p.add_run('C.\tSYOR PEGAWAI KERJA SOSIAL PERUBATAN')
    run.bold = True

    doc.add_paragraph('Berdasarkan siasatan kami, adalah dicadangkan supaya pesakit ini dibantu untuk mendapatkan bantuan...')
    doc.add_paragraph()
    doc.add_paragraph()

    # Tandatangan — jadual 2 kolom supaya selari
    t2 = doc.add_table(rows=0, cols=2)
    t2.autofit = True
    t2.columns[0].width = Cm(7.5)
    t2.columns[1].width = Cm(10.0)
    _add_label_value_row(t2, 'Tandatangan Pegawai', '')
    _add_label_value_row(t2, 'Nama Penuh', '')
    _add_label_value_row(t2, 'Cop Jabatan', '')
    _add_label_value_row(t2, 'Tarikh', data['tarikh_tdi'])

    fpath = os.path.join(output_dir, 'BORANG ULASAN TDI.docx')
    doc.save(fpath)
    return fpath


@app.route('/generate_tdi', methods=['POST'])
def generate_tdi():
    data = prepare_data(request.form)

    # Guna field TDI khusus jika diisi, jika tidak guna dari Seksyen 1
    data['nama'] = request.form.get('tdi_nama', '').strip() or data['nama']
    data['ic'] = request.form.get('tdi_ic', '').strip() or data['ic']
    data['diagnosa'] = request.form.get('tdi_diagnosa', '').strip() or data['diagnosa']

    name = data['nama'].strip()
    if not name:
        return jsonify({'error': 'Sila isi Nama Pesakit dahulu.'}), 400

    # Buat folder atas nama pesakit dalam Dropbox TDI structure
    # Guna nama asal (perkataan asal, ibu huruf besar) supaya padan dengan folder sedia ada
    safe_name = ' '.join(name.strip().split())
    safe_name = safe_name.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
    patient_base = os.path.join(DROPBOX_TDI, safe_name)
    os.makedirs(patient_base, exist_ok=True)

    # Auto-numbering: cari folder nombor seterusnya (1, 2, 3...)
    existing = []
    for item in os.listdir(patient_base):
        sub = os.path.join(patient_base, item)
        if os.path.isdir(sub) and item.isdigit():
            existing.append(int(item))
    next_num = (max(existing) + 1) if existing else 1
    patient_dir = os.path.join(patient_base, str(next_num))
    os.makedirs(patient_dir, exist_ok=True)

    # Auto-tetapkan rujukan jika kosong
    # Prioriti: input manual Rujukan TDI > No. Rujukan Fail (UKSP/HHT/{fail}) > fallback
    rujukan_input = request.form.get('rujukan', '').strip()
    rujukan_fail = request.form.get('rujukan_fail', '').strip()
    if rujukan_input:
        data['rujukan'] = rujukan_input
    elif rujukan_fail:
        data['rujukan'] = f"UKSP/HHT/{rujukan_fail}"
    else:
        data['rujukan'] = f"HHT/UKSP/MSW15/{datetime.now().year}"
    data['tarikh_tdi'] = request.form.get('tarikh_tdi', '').strip() or _format_tarikh(datetime.now())

    # Field TDI
    data['tdi_rawatan'] = request.form.get('tdi_rawatan', '').strip()
    data['tdi_kos'] = request.form.get('tdi_kos', 'Tidak Berkenaan').strip()
    data['tdi_mampu'] = request.form.get('tdi_mampu', 'Tiada').strip()
    data['tdi_ulasan'] = request.form.get('tdi_ulasan', '').strip()
    data['tdi_syor'] = request.form.get('tdi_syor', '').strip()

    try:
        surat_path = generate_surat_iringan(data, patient_dir)
        ulasan_path = generate_borang_ulasan(data, patient_dir)

        return jsonify({
            'ok': True,
            'folder': patient_dir,
            'surat': os.path.basename(surat_path),
            'ulasan': os.path.basename(ulasan_path),
            'file_count': len(os.listdir(patient_dir))
        })
    except Exception as e:
        return jsonify({'error': f'Gagal jana dokumen TDI: {str(e)}'}), 500


DRAFT_FILE = os.path.join(DROPBOX_APP, 'drafts.json')

@app.route('/save_draft', methods=['POST'])
def save_draft():
    data = request.form.to_dict()
    # Gabungkan multi-value fields (checkboxes) jadi string comma-separated
    utiliti_list = request.form.getlist('utiliti')
    if utiliti_list:
        data['utiliti'] = ', '.join(utiliti_list)
    drafts = {}
    if os.path.exists(DRAFT_FILE):
        with open(DRAFT_FILE, 'r') as f:
            drafts = json.load(f)
    key = data.get('draft_key', data.get('nama', 'draft_' + datetime.now().strftime('%Y%m%d_%H%M%S')))
    drafts[key] = data
    with open(DRAFT_FILE, 'w') as f:
        json.dump(drafts, f, indent=2)
    return jsonify({'ok': True, 'key': key})

@app.route('/list_drafts', methods=['GET'])
def list_drafts():
    drafts = {}
    if os.path.exists(DRAFT_FILE):
        with open(DRAFT_FILE, 'r') as f:
            drafts = json.load(f)
    keys = []
    for k, v in drafts.items():
        keys.append({'key': k, 'nama': v.get('nama', 'Tanpa Nama'), 'tarikh': v.get('tarikh_daftar', '')})
    return jsonify(keys)

@app.route('/load_draft/<key>', methods=['GET'])
def load_draft(key):
    if os.path.exists(DRAFT_FILE):
        with open(DRAFT_FILE, 'r') as f:
            drafts = json.load(f)
        data = drafts.get(key, {})
        return jsonify(data)
    return jsonify({})


# ─────────────────────────────────────────────────
# Auto-Update via manifest.json
# ─────────────────────────────────────────────────

# Senarai folder yang selamat untuk diupdate (elak overwrite data)
ALLOWED_UPDATE_PATHS = ('app.py', 'templates/', 'static/')

@app.route('/check_update', methods=['GET'])
def check_update():
    """Bandingkan versi semasa dengan versi dalam manifest."""
    # Dalam EXE, auto-update tak disokong — bagi tahu user.
    if IS_EXE:
        return jsonify({
            'current': APP_VERSION,
            'latest': APP_VERSION,
            'has_update': False,
            'is_exe': True,
            'note': 'Versi EXE — kemaskini melalui pengedar semula (hubungi penghasil app untuk BPS.exe terkini).'
        })
    if not UPDATE_URL:
        return jsonify({
            'error': 'URL update belum dikonfigurasi.',
            'hint': 'Tambah baris UPDATE_URL = <url> dalam config.txt'
        })
    try:
        r = requests.get(UPDATE_URL, timeout=10)
        r.raise_for_status()
        manifest = r.json()
    except requests.exceptions.ConnectionError:
        return jsonify({'error': '❌ Tiada internet — tak dapat capai server update.'})
    except Exception as e:
        return jsonify({'error': f'❌ Gagal baca manifest: {e}'})

    latest = manifest.get('version', '')
    files = manifest.get('files', [])
    change_log = manifest.get('change_log', 'Tiada catatan perubahan.')

    try:
        has_update = version.parse(latest) > version.parse(APP_VERSION)
    except Exception:
        has_update = latest > APP_VERSION  # fallback jika format pelik

    return jsonify({
        'current': APP_VERSION,
        'latest': latest,
        'has_update': has_update,
        'files_count': len(files),
        'files': [f['path'] for f in files],
        'change_log': change_log
    })


@app.route('/apply_update', methods=['POST'])
def apply_update():
    """Muat turun fail terkini dari manifest dan gantikan fail tempatan."""
    if IS_EXE:
        return jsonify({'error': 'Versi EXE tidak menyokong auto-update. '
                                  'Sila dapatkan BPS.exe terkini dari penghasil app.'})
    if not UPDATE_URL:
        return jsonify({'error': 'URL update belum dikonfigurasi.'})
    try:
        r = requests.get(UPDATE_URL, timeout=10)
        manifest = r.json()
    except Exception as e:
        return jsonify({'error': f'Gagal muat turun manifest: {e}'})

    latest = manifest.get('version', '')
    try:
        has_update = version.parse(latest) > version.parse(APP_VERSION)
    except Exception:
        has_update = latest > APP_VERSION
    if not has_update:
        return jsonify({'error': 'Tiada update baru.', 'current': APP_VERSION, 'latest': latest})

    hasil = []
    for f in manifest.get('files', []):
        path = f.get('path', '')
        furl = f.get('url', '')
        if not path or not furl:
            continue

        # Keselamatan: hanya fail dalam ALLOWED_UPDATE_PATHS
        allowed = False
        for prefix in ALLOWED_UPDATE_PATHS:
            if path == prefix or path.startswith(prefix):
                allowed = True
                break
        if not allowed:
            hasil.append(f'⏭️ {path} — dilangkau (bukan dalam senarai selamat)')
            continue

        dest_path = os.path.join(APP_DIR, path)
        dir_name = os.path.dirname(dest_path)

        try:
            # Backup fail lama
            if os.path.exists(dest_path):
                bak_path = dest_path + '.bak'
                with open(dest_path, 'rb') as src, open(bak_path, 'wb') as dst:
                    dst.write(src.read())

            # Muat turun fail baru
            fr = requests.get(furl, timeout=20)
            fr.raise_for_status()
            os.makedirs(dir_name, exist_ok=True)
            with open(dest_path, 'wb') as fh:
                fh.write(fr.content)

            hasil.append(f'✅ {path}')
        except Exception as e:
            hasil.append(f'❌ {path}: {e}')

    return jsonify({
        'result': hasil,
        'latest': latest,
        'restart': True,
        'message': '✅ Update siap! Sila restart app (tutup & buka semula) untuk kesan perubahan.'
    })


# ─────────────────────────────────────────────────
# Jalan
# ─────────────────────────────────────────────────
if __name__ == '__main__':
    import webbrowser
    import threading
    print("=" * 60)
    print(f"  BPS REPORT GENERATOR v{APP_VERSION}")
    print("  Laporan Penilaian Biopsikososial")
    print("=" * 60)
    print(f"  Buka: http://localhost:5000")
    print(f"  Output: {OUTPUT_DIR}")
    print("  Ctrl+C untuk berhenti")
    print("=" * 60)
    # Auto-buka browser lepas server ready (berguna bila jadi .exe)
    def _buka_browser():
        import time
        time.sleep(2)
        webbrowser.open('http://localhost:5000')
    threading.Thread(target=_buka_browser, daemon=True).start()
    app.run(debug=False, host='127.0.0.1', port=5000)
